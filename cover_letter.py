from docx import Document
import re 
import os

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)


def main():
    if not os.path.isdir('./input'):  # If the data directory does not exist, create it
        os.mkdir('./input')

    company = "Barber Collins"
    title = "Junior Copywriter Extraordinaire"
    regex1 = re.compile("\<company\>")
    regex2 = re.compile("\<title\>")

    filename = './input/cover_letter.docx'
    doc = Document(filename)
    docx_replace_regex(doc, regex1 , company)
    docx_replace_regex(doc, regex2 , title)
    doc.save('./input/result.docx')

    # for p in doc.paragraphs:
    #     print(p.text)

if __name__ == "__main__":
    main()